from PyPDF2 import PdfReader
from docx import Document
import os


def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        pdf_reader = PdfReader(file_path)
        print(f"PDF loaded. Total pages: {len(pdf_reader.pages)}")
        
        for page_num, page in enumerate(pdf_reader.pages):
            text += page.extract_text()
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text


def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        print(f"DOCX loaded. Total paragraphs: {len(doc.paragraphs)}")
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text


def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        print(f"TXT loaded. Total characters: {len(text)}")
    except Exception as e:
        print(f"Error reading TXT: {e}")
    return text


def extract_text(file_path):
    """Extract text based on file extension"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def split_into_chunks(text, chunk_size=250, overlap=30):
    """Split text into smaller overlapping chunks"""
    chunks = []
    
    # Split by sentences first
    sentences = text.replace("\n", " ").split(". ")
    
    current_chunk = ""
    for sentence in sentences:
        sentence = sentence.strip() + "."
        
        if len(current_chunk) + len(sentence) < chunk_size:
            current_chunk += " " + sentence
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = sentence
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks
